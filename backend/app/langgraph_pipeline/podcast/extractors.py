import os
import re
import logging
import requests
from io import BytesIO
from bs4 import BeautifulSoup
from docx import Document
import pdfplumber
from typing import List, Optional

logger = logging.getLogger(__name__)

# 옵션: 다운로드 파일을 로컬에 저장할지 여부
SAVE_DOWNLOADED_FILES = False  # True로 설정하면 디버깅 가능

# 프로젝트 루트 기준 출력 폴더 - 절대 경로로 수정
OUTPUT_SAVE_DIR = os.path.abspath("outputs/podcasts")
os.makedirs(OUTPUT_SAVE_DIR, exist_ok=True)

# print(f"[EXTRACTORS] Output directory: {OUTPUT_SAVE_DIR}")

class TextExtractor:
    """다양한 소스에서 텍스트 추출"""

    # -------------------------
    # 공통 - URL 다운로드
    # -------------------------
    @staticmethod
    def download_file(url: str, suffix: str) -> BytesIO:
        """
        URL을 다운로드하여 메모리(BytesIO)로 반환
        필요시 로컬 저장도 가능(SAVE_DOWNLOADED_FILES 사용)
        """
        try:
            logger.info(f"[DOWNLOAD START] URL: {url[:100]}...")
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            
            resp = requests.get(url, timeout=30, headers=headers, stream=True)
            resp.raise_for_status()
            
            # 전체 내용을 메모리에 읽기
            file_bytes = resp.content
            logger.info(f"[DOWNLOAD SUCCESS] Size: {len(file_bytes):,} bytes")

            if SAVE_DOWNLOADED_FILES:
                # URL에서 파일명 추출 (query parameter 제거)
                filename = url.split("?")[0].split("/")[-1]
                
                # 파일명이 너무 길거나 이상하면 기본 이름 사용
                if not filename or len(filename) > 100 or not any(c.isalnum() for c in filename):
                    import time
                    filename = f"download_{int(time.time())}{suffix}"
                
                save_path = os.path.join(OUTPUT_SAVE_DIR, filename)
                
                try:
                    with open(save_path, "wb") as f:
                        f.write(file_bytes)
                    logger.info(f"[FILE SAVED] {save_path}")
                    print(f"✅ 다운로드 파일 저장됨: {save_path}")
                except Exception as save_error:
                    logger.warning(f"[SAVE FAILED] {save_error}")

            return BytesIO(file_bytes)
            
        except requests.exceptions.Timeout:
            logger.error("[DOWNLOAD ERROR] Timeout (30s)")
            raise Exception(f"파일 다운로드 타임아웃: {url[:100]}")
        except requests.exceptions.RequestException as e:
            logger.error(f"[DOWNLOAD ERROR] {e}")
            raise Exception(f"파일 다운로드 실패: {str(e)}")

    # -------------------------
    # 웹 페이지 텍스트 추출
    # -------------------------
    @staticmethod
    def extract_from_web(url: str) -> str:
        try:
            logger.info(f"[WEB EXTRACT] {url[:100]}...")
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            resp = requests.get(url, timeout=10, headers=headers)
            resp.raise_for_status()
            soup = BeautifulSoup(resp.text, 'html.parser')

            for tag in soup(['script', 'style']):
                tag.decompose()

            text = soup.get_text(separator="\n", strip=True)
            cleaned = re.sub(r"\n{3,}", "\n\n", text)
            logger.info(f"[WEB SUCCESS] Extracted {len(cleaned)} characters")
            return cleaned
        except Exception as e:
            logger.error(f"[WEB ERROR] {e}")
            return ""

    # -------------------------
    # DOCX 추출
    # -------------------------
    @staticmethod
    def extract_from_docx(source: str) -> str:
        """
        source가 URL이면 다운로드하여 메모리에서 처리
        source가 로컬 경로면 그대로 처리
        """
        try:
            logger.info(f"[DOCX EXTRACT START] Source: {source[:100]}...")
            
            if source.startswith("http"):
                file_obj = TextExtractor.download_file(source, suffix=".docx")
                doc = Document(file_obj)
            else:
                if not os.path.exists(source):
                    raise FileNotFoundError(f"DOCX 파일을 찾을 수 없습니다: {source}")
                doc = Document(source)

            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            logger.info(f"[DOCX SUCCESS] Extracted {len(text)} characters")
            
            # 미리보기 출력
            if text:
                preview = text[:200].replace("\n", " ")
                print(f"📄 DOCX Preview: {preview}...")
            
            return text

        except Exception as e:
            logger.error(f"[DOCX ERROR] {e}", exc_info=True)
            raise Exception(f"DOCX 추출 실패: {str(e)}")

    # -------------------------
    # PDF 추출
    # -------------------------
    @staticmethod
    def extract_from_pdf(source: str) -> str:
        """
        source가 URL → 메모리에서 pdfplumber 처리
        source가 로컬 경로 → 기존 처리
        """
        pdf_stream = None
        try:
            logger.info(f"[PDF EXTRACT START] Source: {source[:100]}...")
            
            if source.startswith("http"):
                file_obj = TextExtractor.download_file(source, suffix=".pdf")
                pdf_stream = file_obj
            else:
                if not os.path.exists(source):
                    raise FileNotFoundError(f"PDF 파일을 찾을 수 없습니다: {source}")
                pdf_stream = open(source, "rb")

            text = ""
            with pdfplumber.open(pdf_stream) as pdf:
                total_pages = len(pdf.pages)
                logger.info(f"[PDF INFO] Total pages: {total_pages}")
                
                for idx, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"
                    
                    # 첫 페이지 미리보기
                    if idx == 0 and page_text:
                        preview = page_text[:200].replace("\n", " ")
                        logger.info(f"[PDF PREVIEW] First page: {preview}...")
                        print(f"📕 PDF Preview (Page 1): {preview}...")

            # 유니코드 제어 문자만 제거
            text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
            
            logger.info(f"[PDF SUCCESS] Extracted {len(text)} characters from {total_pages} pages")
            
            if not text or len(text.strip()) < 10:
                logger.warning("[PDF WARNING] Extracted text is too short or empty")
                print("⚠️  경고: PDF에서 추출된 텍스트가 없거나 매우 짧습니다.")
            
            return text.strip()

        except Exception as e:
            logger.error(f"[PDF ERROR] {e}", exc_info=True)
            raise Exception(f"PDF 추출 실패: {str(e)}")

        finally:
            if pdf_stream and not source.startswith("http"):
                try:
                    pdf_stream.close()
                except:
                    pass

    # -------------------------
    # 자동 타입 판별
    # -------------------------
    @classmethod
    def extract(cls, source: str) -> str:
        s = source.strip()
        
        logger.info(f"[EXTRACT] Detecting type for: {s[:100]}...")

        # 1) 웹 링크
        if s.startswith("http://") or s.startswith("https://"):
            # URL에서 확장자 추출 (query parameter 제거)
            url_path = s.split("?")[0].lower()
            
            if url_path.endswith(".pdf"):
                logger.info("[TYPE] PDF URL detected")
                return cls.extract_from_pdf(s)
            elif url_path.endswith(".docx"):
                logger.info("[TYPE] DOCX URL detected")
                return cls.extract_from_docx(s)
            else:
                logger.info("[TYPE] Web page detected")
                return cls.extract_from_web(s)

        # 2) 로컬 파일
        if s.lower().endswith(".pdf"):
            logger.info("[TYPE] Local PDF file")
            return cls.extract_from_pdf(s)

        if s.lower().endswith(".docx"):
            logger.info("[TYPE] Local DOCX file")
            return cls.extract_from_docx(s)

        raise ValueError(f"지원하지 않는 소스 타입: {source}")


# -------------------------------------
# 여러 소스 일괄 처리
# -------------------------------------
def extract_all_sources(sources: List[str]) -> tuple[List[str], List[str]]:
    extracted = []
    errors = []

    print(f"\n{'='*80}")
    print(f"📚 텍스트 추출 시작: {len(sources)}개 소스")
    print(f"{'='*80}\n")

    for i, src in enumerate(sources):
        name = os.path.basename(src) if not src.startswith("http") else src[:50]
        logger.info(f"[EXTRACT] {i+1}/{len(sources)} → {name}")
        print(f"\n[{i+1}/{len(sources)}] 처리 중: {name}...")

        try:
            text = TextExtractor.extract(src)

            if text and len(text.strip()) > 0:
                extracted.append(text)
                print(f"✅ 성공: {len(text):,} 글자 추출됨")
                logger.info(f"[SUCCESS] Extracted {len(text)} characters")
            else:
                error_msg = f"{name}: 텍스트 추출 실패 (빈 결과)"
                errors.append(error_msg)
                print(f"❌ 실패: {error_msg}")
                logger.warning(error_msg)
                
        except Exception as e:
            error_msg = f"{name}: 처리 오류 → {str(e)}"
            errors.append(error_msg)
            print(f"❌ 실패: {error_msg}")
            logger.error(error_msg, exc_info=True)

    print(f"\n{'='*80}")
    print(f"📊 추출 완료: 성공 {len(extracted)}개 / 실패 {len(errors)}개")
    print(f"{'='*80}\n")
    
    logger.info(f"[EXTRACT COMPLETE] Success: {len(extracted)}, Errors: {len(errors)}")
    return extracted, errors